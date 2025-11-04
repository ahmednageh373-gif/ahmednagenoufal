#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=======================================================================
وحدة توليد التقارير الذكية (Smart Reports Generator)
=======================================================================

الهدف: تجميع مخرجات وحدات التحليل الأخرى في تقرير نهائي موحد ومنسق.

المدخلات:
    - مخرجات quantity_analysis.py
    - مخرجات schedule_analysis.py
    - صور منحنيات S-Curve

المخرجات:
    - تقرير نهائي منسق (DOCX/PDF)
    - تقرير تحليلي ذكي باللغة العربية

الممارسات الهندسية المدمجة:
    - التنسيق الموحد للتقارير
    - الربط بالرسومات الهندسية
    - تنبيهات الامتثال لكود البناء السعودي
    
التاريخ: 2025-11-04
المطور: NOUFAL Engineering System
=======================================================================
"""

import pandas as pd
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

# إعداد السجلات
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class SmartReportGenerator:
    """
    مولد التقارير الذكية
    Smart Report Generator
    """
    
    def __init__(self, project_path: str, project_name: str = "مشروع هندسي"):
        """
        تهيئة المولد
        
        Args:
            project_path: المسار الرئيسي للمشروع
            project_name: اسم المشروع
        """
        self.project_path = Path(project_path)
        self.project_name = project_name
        
        # Input paths
        self.quantities_path = self.project_path / "03_Output_Data" / "01_Quantities"
        self.visuals_path = self.project_path / "03_Output_Data" / "03_Visuals" / "S_Curves"
        self.temp_path = self.project_path / "02_Processing" / "Temp_Data"
        
        # Output path
        self.output_path = self.project_path / "03_Output_Data" / "02_Reports" / "Smart_Reports"
        self.output_path.mkdir(parents=True, exist_ok=True)
        
        # Data containers
        self.quantity_data = None
        self.schedule_data = None
        self.validation_results = []
        
        logger.info(f"تم تهيئة مولد التقارير الذكية للمشروع: {project_name}")
    
    def load_quantity_data(self, filename: str = "Final_BOQ.xlsx") -> Dict:
        """
        تحميل بيانات الكميات
        
        Args:
            filename: اسم ملف الكميات
            
        Returns:
            Dict: بيانات الكميات
        """
        try:
            file_path = self.quantities_path / "Final_BOQ" / filename
            logger.info(f"تحميل بيانات الكميات من: {file_path}")
            
            # Read BOQ data
            df_boq = pd.read_excel(file_path, sheet_name='BOQ')
            
            # Read validation results if exists
            try:
                df_validation = pd.read_excel(file_path, sheet_name='SBC Compliance')
                validation_results = df_validation.to_dict('records')
            except:
                validation_results = []
            
            self.quantity_data = {
                'boq': df_boq,
                'validation': validation_results,
                'total_items': len(df_boq),
                'total_quantity': df_boq['calculated_quantity'].sum() if 'calculated_quantity' in df_boq.columns else 0
            }
            
            self.validation_results = validation_results
            
            logger.info(f"تم تحميل {len(df_boq)} بند كميات")
            return self.quantity_data
            
        except Exception as e:
            logger.error(f"خطأ في تحميل بيانات الكميات: {e}")
            raise
    
    def load_schedule_data(self, results: Dict = None) -> Dict:
        """
        تحميل بيانات الجدول الزمني
        
        Args:
            results: نتائج تحليل الجدول (من schedule_analysis.py)
            
        Returns:
            Dict: بيانات الجدول
        """
        self.schedule_data = results
        logger.info("تم تحميل بيانات الجدول الزمني")
        return self.schedule_data
    
    def generate_analytical_text(self) -> str:
        """
        توليد النص التحليلي الذكي
        
        Returns:
            str: النص التحليلي
        """
        logger.info("بدء توليد النص التحليلي...")
        
        analysis = []
        
        # Executive Summary
        analysis.append("ملخص تنفيذي")
        analysis.append("=" * 50)
        analysis.append(f"المشروع: {self.project_name}")
        analysis.append(f"تاريخ التقرير: {datetime.now().strftime('%Y-%m-%d')}")
        analysis.append("")
        
        # Quantity Analysis
        if self.quantity_data:
            analysis.append("📊 تحليل الكميات")
            analysis.append("-" * 50)
            analysis.append(f"• إجمالي البنود: {self.quantity_data['total_items']}")
            analysis.append(f"• إجمالي الكميات: {self.quantity_data['total_quantity']:.2f}")
            
            # SBC Compliance
            critical_issues = len([v for v in self.validation_results if v.get('severity') == 'critical'])
            warnings = len([v for v in self.validation_results if v.get('severity') == 'warning'])
            
            if critical_issues > 0:
                analysis.append(f"⚠️  تحذير: تم رصد {critical_issues} مشكلة حرجة في الامتثال لكود البناء السعودي")
            elif warnings > 0:
                analysis.append(f"ℹ️  ملاحظة: تم رصد {warnings} تحذير في الامتثال لكود البناء السعودي")
            else:
                analysis.append("✅ جميع البنود متوافقة مع كود البناء السعودي (SBC)")
            
            analysis.append("")
        
        # Schedule Analysis
        if self.schedule_data:
            analysis.append("📅 تحليل الجدول الزمني")
            analysis.append("-" * 50)
            analysis.append(f"• إجمالي الأنشطة: {self.schedule_data['total_activities']}")
            analysis.append(f"• الأنشطة الحرجة: {self.schedule_data['critical_activities']}")
            analysis.append(f"• التقدم المخطط: {self.schedule_data['avg_planned_progress']:.1f}%")
            analysis.append(f"• التقدم الفعلي: {self.schedule_data['avg_actual_progress']:.1f}%")
            analysis.append("")
            
            # Status breakdown
            analysis.append("حالة الأنشطة:")
            analysis.append(f"  ✅ على المسار: {self.schedule_data['on_track_activities']} نشاط")
            analysis.append(f"  🚀 متقدم: {self.schedule_data['ahead_activities']} نشاط")
            analysis.append(f"  ⏰ متأخر: {self.schedule_data['behind_activities']} نشاط")
            analysis.append("")
            
            # EVM Analysis
            if 'evm_metrics' in self.schedule_data:
                evm = self.schedule_data['evm_metrics']
                analysis.append("💰 تحليل القيمة المكتسبة (EVM)")
                analysis.append("-" * 50)
                
                # CPI Analysis
                cpi = evm['cost_performance_index']
                if cpi > 1.0:
                    analysis.append(f"✅ مؤشر أداء التكلفة (CPI): {cpi:.2f} - أداء ممتاز (أقل من الموازنة)")
                elif cpi >= 0.9:
                    analysis.append(f"ℹ️  مؤشر أداء التكلفة (CPI): {cpi:.2f} - أداء مقبول")
                else:
                    analysis.append(f"⚠️  مؤشر أداء التكلفة (CPI): {cpi:.2f} - تحذير: تجاوز التكلفة")
                
                # SPI Analysis
                spi = evm['schedule_performance_index']
                if spi > 1.0:
                    analysis.append(f"✅ مؤشر أداء الجدول (SPI): {spi:.2f} - متقدم عن الجدول")
                elif spi >= 0.9:
                    analysis.append(f"ℹ️  مؤشر أداء الجدول (SPI): {spi:.2f} - على المسار")
                else:
                    analysis.append(f"⚠️  مؤشر أداء الجدول (SPI): {spi:.2f} - تحذير: تأخير في الجدول")
                
                analysis.append("")
                analysis.append(f"• القيمة المكتسبة (EV): {evm['earned_value']:,.2f} ريال")
                analysis.append(f"• التكلفة الفعلية (AC): {evm['actual_cost']:,.2f} ريال")
                analysis.append(f"• التقدير عند الإنجاز (EAC): {evm['estimate_at_completion']:,.2f} ريال")
                analysis.append(f"• الانحراف عند الإنجاز (VAC): {evm['variance_at_completion']:,.2f} ريال")
                analysis.append("")
        
        # Recommendations
        analysis.append("🎯 التوصيات")
        analysis.append("-" * 50)
        
        if self.validation_results:
            critical_items = [v for v in self.validation_results if v.get('severity') == 'critical']
            if critical_items:
                analysis.append("1. معالجة القضايا الحرجة في الامتثال لكود البناء السعودي:")
                for item in critical_items[:3]:  # Top 3
                    analysis.append(f"   • {item.get('description')}: {item.get('issue')}")
                analysis.append("")
        
        if self.schedule_data and self.schedule_data['behind_activities'] > 0:
            analysis.append("2. التركيز على الأنشطة المتأخرة لإعادة المشروع إلى المسار الصحيح")
            analysis.append("")
        
        if self.schedule_data and self.schedule_data.get('evm_metrics', {}).get('cost_performance_index', 1.0) < 0.9:
            analysis.append("3. مراجعة التكاليف واتخاذ إجراءات تصحيحية لتحسين مؤشر CPI")
            analysis.append("")
        
        analysis.append("4. متابعة الأنشطة الحرجة بشكل يومي لضمان عدم التأخير")
        analysis.append("5. الالتزام بمعايير كود البناء السعودي في جميع أعمال التنفيذ")
        
        return "\n".join(analysis)
    
    def create_word_report(self, output_filename: str = None) -> str:
        """
        إنشاء التقرير بصيغة Word
        
        Args:
            output_filename: اسم الملف المخرج (اختياري)
            
        Returns:
            str: مسار الملف المحفوظ
        """
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"Smart_Report_{timestamp}.docx"
        
        output_file = self.output_path / output_filename
        
        logger.info(f"بدء إنشاء تقرير Word: {output_file}")
        
        # Create document
        doc = Document()
        
        # Set RTL for Arabic
        self._set_rtl(doc)
        
        # Title
        title = doc.add_heading(f'تقرير ذكي - {self.project_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f'تاريخ التقرير: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Add analytical text
        analysis_text = self.generate_analytical_text()
        
        for line in analysis_text.split('\n'):
            if line.startswith('='):
                doc.add_heading(line.replace('=', '').strip(), level=1)
            elif line.startswith('-'):
                doc.add_heading(line.replace('-', '').strip(), level=2)
            elif line.strip():
                p = doc.add_paragraph(line)
                
                # Color coding
                if '⚠️' in line or 'تحذير' in line:
                    self._set_paragraph_color(p, RGBColor(255, 0, 0))
                elif '✅' in line or 'ممتاز' in line:
                    self._set_paragraph_color(p, RGBColor(0, 128, 0))
        
        # Add page break
        doc.add_page_break()
        
        # Add S-Curve if available
        scurve_files = list(self.visuals_path.glob("*.png"))
        if scurve_files:
            doc.add_heading('منحنى S-Curve', level=1)
            
            # Add most recent S-curve
            latest_scurve = max(scurve_files, key=lambda p: p.stat().st_mtime)
            doc.add_picture(str(latest_scurve), width=Inches(6))
            
            doc.add_paragraph()
        
        # Add SBC Compliance Details
        if self.validation_results:
            doc.add_page_break()
            doc.add_heading('تفاصيل الامتثال لكود البناء السعودي', level=1)
            
            # Create table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = 'رقم البند'
            header_cells[1].text = 'الوصف'
            header_cells[2].text = 'المشكلة'
            header_cells[3].text = 'الخطورة'
            
            # Data rows
            for validation in self.validation_results:
                row_cells = table.add_row().cells
                row_cells[0].text = str(validation.get('item_number', ''))
                row_cells[1].text = str(validation.get('description', ''))
                row_cells[2].text = str(validation.get('issue', ''))
                row_cells[3].text = str(validation.get('severity', ''))
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph('تم إنشاء هذا التقرير بواسطة:')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_para2 = doc.add_paragraph('NOUFAL Engineering Management System')
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para2.runs[0]
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(54, 96, 146)
        
        # Save document
        doc.save(output_file)
        
        logger.info(f"تم حفظ التقرير: {output_file}")
        return str(output_file)
    
    def _set_rtl(self, doc):
        """Set document to RTL (Right-to-Left) for Arabic"""
        sections = doc.sections
        for section in sections:
            section_properties = section._sectPr
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            section_properties.append(bidi)
    
    def _set_paragraph_color(self, paragraph, color: RGBColor):
        """Set paragraph text color"""
        for run in paragraph.runs:
            run.font.color.rgb = color
    
    def run_full_report_generation(self,
                                   quantity_file: str = "Final_BOQ.xlsx",
                                   schedule_results: Dict = None,
                                   output_filename: str = None) -> Dict:
        """
        تشغيل التوليد الكامل للتقرير
        
        Args:
            quantity_file: اسم ملف الكميات
            schedule_results: نتائج تحليل الجدول
            output_filename: اسم الملف المخرج (اختياري)
            
        Returns:
            Dict: نتائج التوليد
        """
        logger.info("=" * 70)
        logger.info("بدء توليد التقرير الذكي")
        logger.info("=" * 70)
        
        # Step 1: Load quantity data
        self.load_quantity_data(quantity_file)
        
        # Step 2: Load schedule data
        if schedule_results:
            self.load_schedule_data(schedule_results)
        
        # Step 3: Generate analytical text
        analysis_text = self.generate_analytical_text()
        
        # Step 4: Create Word report
        report_path = self.create_word_report(output_filename)
        
        # Prepare results
        results = {
            'success': True,
            'report_file': report_path,
            'analysis_text': analysis_text,
            'has_quantity_data': self.quantity_data is not None,
            'has_schedule_data': self.schedule_data is not None,
            'validation_issues': len(self.validation_results)
        }
        
        logger.info("=" * 70)
        logger.info("اكتمل توليد التقرير بنجاح!")
        logger.info(f"ملف التقرير: {report_path}")
        logger.info(f"قضايا الامتثال: {results['validation_issues']}")
        logger.info("=" * 70)
        
        return results


def main():
    """
    دالة الاختبار الرئيسية
    """
    project_path = "/home/user/webapp/backend/project_template"
    project_name = "مشروع فيلا سكنية"
    
    generator = SmartReportGenerator(project_path, project_name)
    
    print("=" * 70)
    print("وحدة توليد التقارير الذكية - NOUFAL Engineering System")
    print("=" * 70)
    print("\nالوحدة جاهزة للاستخدام!")
    print("\nمثال على الاستخدام:")
    print("  generator = SmartReportGenerator('/path/to/project', 'Project Name')")
    print("  results = generator.run_full_report_generation()")
    print("=" * 70)


if __name__ == "__main__":
    main()
